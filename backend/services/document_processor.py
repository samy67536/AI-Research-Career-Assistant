# backend/services/document_processor.py
"""
Document Processing Service
Handles: PDF extraction (PyPDF + PDFPlumber), DOCX extraction, text cleaning, chunking
"""
import os
import re
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from config import settings
from typing import List, Dict, Tuple


class DocumentProcessor:

    def __init__(self):
        self.chunk_size    = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap

    # ── Public API ────────────────────────────────────────────────────

    def extract_text(self, file_path: str, file_type: str) -> Tuple[str, Dict]:
        """Extract text and metadata from PDF or DOCX."""
        if file_type == "pdf":
            return self._extract_pdf(file_path)
        elif file_type == "docx":
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def chunk_text(self, text: str, metadata: Dict = None) -> List[Dict]:
        """
        Split text into overlapping chunks using recursive strategy.
        Returns list of dicts: {text, chunk_index, char_start, char_end, metadata}
        """
        separators = ["\n\n", "\n", ". ", " ", ""]
        chunks = self._recursive_split(text, separators)
        result = []
        for i, chunk in enumerate(chunks):
            clean = chunk.strip()
            if len(clean) > 50:   # skip trivially short chunks
                result.append({
                    "text":        clean,
                    "chunk_index": i,
                    "word_count":  len(clean.split()),
                    "metadata":    metadata or {}
                })
        return result

    # ── PDF Extraction ────────────────────────────────────────────────

    def _extract_pdf(self, file_path: str) -> Tuple[str, Dict]:
        text = ""
        page_count = 0

        # Try PDFPlumber first (better layout handling)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception:
            pass

        # Fallback to PyPDF if PDFPlumber yielded nothing
        if not text.strip():
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                page_count = len(reader.pages)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            except Exception as e:
                raise RuntimeError(f"Failed to extract PDF: {e}")

        text = self._clean_text(text)
        metadata = {
            "page_count": page_count,
            "word_count": len(text.split()),
            "file_type":  "pdf",
            "file_path":  file_path
        }
        return text, metadata

    # ── DOCX Extraction ───────────────────────────────────────────────

    def _extract_docx(self, file_path: str) -> Tuple[str, Dict]:
        from docx import Document
        doc   = Document(file_path)
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts)
        text = self._clean_text(text)
        metadata = {
            "page_count": len(doc.sections),
            "word_count": len(text.split()),
            "file_type":  "docx",
            "file_path":  file_path
        }
        return text, metadata

    # ── Text Cleaning ─────────────────────────────────────────────────

    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)           # collapse whitespace
        text = re.sub(r'[\x00-\x08\x0b\x0e-\x1f]', '', text)  # remove control chars
        text = re.sub(r'\.{3,}', '...', text)      # normalize ellipsis
        text = text.strip()
        return text

    # ── Recursive Text Splitter ───────────────────────────────────────

    def _recursive_split(self, text: str, separators: List[str]) -> List[str]:
        if not separators:
            return [text]

        sep = separators[0]
        splits = text.split(sep) if sep else list(text)
        chunks, current = [], ""

        for split in splits:
            piece = (current + sep + split).strip() if current else split.strip()
            if len(piece) <= self.chunk_size:
                current = piece
            else:
                if current:
                    chunks.append(current)
                if len(split) > self.chunk_size:
                    sub_chunks = self._recursive_split(split, separators[1:])
                    chunks.extend(sub_chunks)
                    current = ""
                else:
                    current = split

        if current:
            chunks.append(current)

        # Apply overlap
        overlapped = []
        for i, chunk in enumerate(chunks):
            if i > 0 and self.chunk_overlap > 0:
                prev_words = chunks[i - 1].split()
                overlap_text = " ".join(prev_words[-self.chunk_overlap // 10:])
                chunk = overlap_text + " " + chunk
            overlapped.append(chunk[:self.chunk_size + self.chunk_overlap])

        return overlapped
